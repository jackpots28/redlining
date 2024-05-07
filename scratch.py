# Using python-docx for editing
from typing import Dict, Any, Type

import docx
from docx.document import Document
from docx.enum.text import WD_COLOR_INDEX

from pathlib import Path
import os, sys, string
from src.logger.logger import func_log

test_doc_path = Path("./word_doc_test.docx")
output_doc_path = Path("./output_files/word_doc_test.docx")

print(f"Is the test file path correct: {test_doc_path.is_file()}\n")
print(f"docx version: {docx.__version__}", end="\n\n")

doc = docx.Document(str(test_doc_path))
output_doc = docx.Document()

'''
This is a good working test for stripping paragraphs into dictionaries of lists of strings
'''
def doc_to_dict(input_doc: Document) -> dict[int, list]:
    return {k:[[word.strip(string.punctuation)] \
        for word in (v.text).split(". ") \
            if word] for (k,v) in enumerate(input_doc.paragraphs)}


###
# def extract_words_using_find(input_string: str) -> list[list]:
#     words = [[word.strip(string.punctuation)] for word in input_string.split(". ") if word]
#     return words

# print(f"{ {k:v for (k,v) in doc_to_dict(doc).items()} }")

test_dict = doc_to_dict(doc)

for k, v in test_dict.items():
    #print(f'{k}: {v}')
    for item in v:
        if "ipsum" in str(item):
            print(f"{k}: {item}")

        # temp_list = extract_words_using_find(v)
        # print(temp_list)
        # output_doc.add_paragraph(v)


# output_doc.save("./output_files/test_doc.docx")





# Can Highlight text
# for par in doc.paragraphs:
#     for run in par.runs:
#         if 'ultricies' in run.text:
#             x = run.text.split('ultricies')
#             run.clear()
#             for i in range(len(x)):
#                 print(x[i])
#                 print("CHANGED")
#                 # run.add_text(x[i])
#                 # run.add_text('ultricies')
#                 # run.font.highlight_color = WD_COLOR_INDEX.YELLOW
#                 # run.font.strike = True

# Remove full paragraphs
# def delete_paragraph(paragraph):
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None

# Create list with single word changed
# for paragraph in doc.paragraphs:
#     if 'lorem' in paragraph.text:
#         for run in paragraph.runs:
#             if 'lorem' in run.text:
#                 x = run.text.split('lorem')
#                 run.clear()
#                 # print(x)
#                 for i in range(len(x)-1):
#                     run.add_text(x[i])
#                     run.add_text('lorem')
#                     run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # new_list.append((par.text).replace("ultricies", "CHANGED"))

# for item in new_list:
#     print(item)

# Clear out working doc
# for par in doc.paragraphs:
#     delete_paragraph(par)

# Fill it with new list of paragraphs with the modifications
# for item in new_list:
#     doc.add_paragraph().add_run(str(item))

# Save it to a new file
# doc.save(output_doc_path)

# @func_log
# def doc_to_list(input_doc: docx.Document):
#     paragraph_lists = [[]]
#     for paragraph in input_doc.paragraphs:
#         paragraph_lists.append(paragraph.text)
#
#     return paragraph_lists


# print(doc_to_list(doc))

# dict_of_para = {}
# for i, paragraph in enumerate(doc.paragraphs):
#     dict_of_para[i] = paragraph.text
#     #print(f"{i}: {paragraph.text}\n")
#
#
# for key, value in dict_of_para.items():
#     print(f"{key}: {value}")


# def doc_to_dict(input_doc: docx.Document) -> dict[int, str]:
#     return {k:v.text for (k,v) in enumerate(input_doc.paragraphs)}
#
#
# test_dict = doc_to_dict(doc)
# for key, value in doc_to_dict(doc).items():
#     print(f"{key}: {value}")
# @func_log
# def test_log_wrapper(x, y):
#     return x+y
#
# result = test_log_wrapper(10, 20)

# NOTES
# import docx

# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)