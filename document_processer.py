# Using flask as a frontend for uploading file from localhost chosen directory
from flask import Flask, request, redirect, url_for, send_from_directory,render_template
from werkzeug.utils import secure_filename
import os

import docx

import highlight_word_test

UPLOAD_FOLDER = os.path.abspath("./uploads/")
ALLOWED_EXTENSIONS = set(["txt", "text", "docx","pdf"])

words = ['lorem']

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

@app.route("/")
def index():
    return render_template("upload.html")


@app.route("/upload", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        if not "file" in request.files:
            return "No file part in the form."
        f = request.files["file"]
        if f.filename == "":
            return "No file selected."
        if f and allowed_file(f.filename):
            filename = secure_filename(f.filename)
            doc=docx.Document(filename)
            for word in words:
                doc = highlight_word_test.split_Runs(doc, word)
            for word in words:
                doc = highlight_word_test.style_Token(doc,word,True)

            doc.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
            # f.save(os.path.join(app.config["UPLOAD_FOLDER"], filename))
            return "file successfully upload"
        return "File not allowed."
    return "Upload file route"


if __name__ == "__main__":
    app.run(debug=True)

# NOTES
# import docx

# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)