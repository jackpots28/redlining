import docx
import re
from docx.enum.text import WD_COLOR_INDEX

from pathlib import Path
import os

#TODO - THIS NEEDS TO BE DONE AS A CLASS STRUCTURE TO BE CALLED EXTERNALLY IN OTHER SRC FILES


def test_split_sentence_to_list(text: str) -> list:
    output_list = list()
    output_list = text.split(" ")
    print(output_list)
    return output_list

def split_runs(doc: docx.Document, word: str) -> docx.Document:
    for p in doc.paragraphs:
        print(p.text.find(word))
        if p.text.find(word) != -1:
            DEBUGGING_OUTPUT_TEXT = p.text
            virtualRuns=p.runs
            p.text = ""
            for r in virtualRuns:
                if r.text.find(word) != -1:
                    broken_sentence = test_split_sentence_to_list(r.text)
                    for word in broken_sentence:
                        p.add_run(word)
                        p.add_run(" ")
    return doc
    
def style_token(doc: docx.Document, word: str, comment=True) -> docx.Document:
    for p in doc.paragraphs:
        for i,r in enumerate(p.runs):
            if p.runs[i].text.find(word) != -1:
                p.runs[i].font.highlight_color = WD_COLOR_INDEX.RED
                p.runs[i].font.strike=True
    return doc


# words = ['ipsum']

# for word in words:
#     doc = split_Runs(doc, word)

# for word in words:
#     doc = style_Token(doc,word,True)

#nums is the list of tokens that is going to be highlighted and comment
# nums=['vulputate ']
# for num in nums:
#     doc=split_Runs(doc,num)
# for num in nums:
#     doc=style_Token(doc,num,True)

# doc.save(output_doc_path)