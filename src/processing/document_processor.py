from typing import Dict, Any, List

from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path
from docx.document import Document

import docx
import os, sys, string


# Setup root_path for referencing files/dirs consistently
project_root = os.path.realpath(os.path.join(os.path.dirname(__file__), '../..'))
sys.path.insert(0, os.path.abspath(project_root))

# Sourcing internal packages
from src.logger import logger
from src.logger.logger import func_log

# Setup Logging
log_path = Path(f"{project_root}/logs/document_processing.log")
document_processing_logger = logger.setup_logger("document_processing_logger", log_path)

# Creates k:v pair where the k is just an index and the value is the entire str found via doc.paragraph
@func_log
def doc_to_dict(input_doc: Document) -> dict[Any, list[list[Any]]]:
    return {k:[[word.strip(string.punctuation)] \
        for word in v.text.split(". ") \
            if word] for (k,v) in enumerate(input_doc.paragraphs)}

@func_log
def split_sentence_to_list(text: str) -> list:
    document_processing_logger.debug(f'Split sentence buffer: {text.split(". ")}')
    return text.split(". ")


# TODO - Need to rewrite the splits to utilize the doc_to_dict, 
# will need to split on ". " = sentence and attempt highlighting full sentence where keyword is found

# Creates a set of "virtual" runs that the "style_token" func loops through
# to apply formatting based on if keyword is found in list of runs
@func_log
def split_runs(doc: Document, word: str) -> Document:
    for p in doc.paragraphs:
        # document_processing_logger.debug(f"Boolean value if word is found: {p.text.find(word)}")
        if word in p.text:
            # DEBUGGING_OUTPUT_TEXT = p.text
            # document_processing_logger.debug(f"Text buffer as of found: {word}:\n{DEBUGGING_OUTPUT_TEXT}")
            virtual_runs = p.runs
            p.text = ""
            for r in virtual_runs:
                if r.text.find(word) != -1:
                    broken_sentence = split_sentence_to_list(r.text)
                    for word in broken_sentence:
                        p.add_run(word)
                        p.add_run(" ")
    return doc


# Loops over doc "runs" and applies text formatting only to "keywords" that are found
@func_log
def style_token(doc: Document, word: str, comment=True) -> Document:
    for p in doc.paragraphs:
        for i, r in enumerate(p.runs):
            document_processing_logger.debug(f"r value in enumeration of runs: {r}")
            if p.runs[i].text.find(word) != -1:
                p.runs[i].font.highlight_color = WD_COLOR_INDEX.RED
                p.runs[i].font.strike = True
    return doc

###

# words = ['ipsum']

# for word in words:
#     doc = split_Runs(doc, word)

# for word in words:
#     doc = style_Token(doc,word,True)

#nums is the list of tokens that is going to be highlighted and comment
# nums=['vulputate ']
# for num in nums:
#     doc=split_Runs(doc,num)
# for num in nums:
#     doc=style_Token(doc,num,True)

# doc.save(output_doc_path)
