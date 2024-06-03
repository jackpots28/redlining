'''
THIS SRATCH WORKS FOR HIGHLIGHTING SENTENCE WITH SPECIFIC WORD
'''

import docx
from docx.document import Document
from docx.enum.text import WD_COLOR_INDEX

import os, sys
from pathlib import Path

project_root = os.path.realpath(os.path.join(os.path.dirname(__file__), '../../..'))
sys.path.insert(0, os.path.abspath(project_root))

test_doc_path = Path("../word_doc_test.docx")
test_doc = docx.Document(test_doc_path)
output_doc_path = Path("./output_files/test_doc.docx")
output_doc = docx.Document(output_doc_path)

word_to_highlight = "faucibus"


def split_runs(doc: Document, word: str) -> Document:
    for p in doc.paragraphs:
        if word in p.text:
            virtual_runs = p.runs
            p.text = ""
            for r in virtual_runs:
                start_index = r.text.find(word)
                if start_index != -1:
                    broken_sentence = r.text.split(". ")
                    for s_word in broken_sentence:
                        p.add_run(s_word)
                        p.add_run(". ")
    return doc


# Loops over doc "runs" and applies text formatting only to "keywords" that are found
def style_token(doc: Document, word: str, comment=True) -> Document:
    for p in doc.paragraphs:
        for i, r in enumerate(p.runs):
            if p.runs[i].text.find(word) != -1:
                p.runs[i].font.highlight_color = WD_COLOR_INDEX.RED
                p.runs[i].font.strike = True
    return doc


doc = split_runs(test_doc, word_to_highlight)
doc = style_token(test_doc, word_to_highlight, True)

doc.save(str(output_doc_path))