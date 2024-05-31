# Using python-docx for editing
from typing import Dict, Any, Type

import docx
from docx.document import Document
from docx.enum.text import WD_COLOR_INDEX

from pathlib import Path
import os, sys, string, tempfile, itertools, re
from src.logger.logger import func_log

project_root = os.path.realpath(os.path.join(os.path.dirname(__file__), '../..'))
sys.path.insert(0, os.path.abspath(project_root))

from src.processing import document_processor

test_doc_path = Path("./word_doc_test.docx")
output_doc_path = Path("./output_files/test_doc.docx")

print(f"Is the test file path correct: {test_doc_path.is_file()}\n")
print(f"docx version: {docx.__version__}", end="\n\n")

doc = docx.Document(str(test_doc_path))
output_doc = docx.Document()

'''
This is a good working test for stripping paragraphs into dictionaries of lists of strings
'''
# def doc_to_dict(input_doc: Document) -> dict[int, list[str]]:
#     return {k:[[word.strip(string.punctuation)] \
#         for word in (v.text).split(". ") \
#             if word] for (k,v) in enumerate(input_doc.paragraphs)}


# def doc_to_dict(input_doc: Document) -> dict[int, str]:
#     return {k:str(word.text).strip(string.punctuation) for (k, word) in enumerate(input_doc.paragraphs)}


def doc_to_dict(input_doc: Document) -> dict[int, Document.paragraphs]:
    return {k:word for (k, word) in enumerate(input_doc.paragraphs)}

'''
TODO - Build dictionary of lists in this format {key, [paragraph, run.formatting.info]}
'''
# def doc_to_dict_w_runs(input_doc: Document) -> dict[int, list[str]]:
#     return {k:[[word.strip(string.punctuation)] \
#         for word in (v.text).split(". ") \
#             if word] for (k,v) in enumerate(input_doc.paragraphs)}

###
# def extract_words_using_find(input_string: str) -> list[list]:
#     words = [[word.strip(string.punctuation)] for word in input_string.split(". ") if word]
#     return words

# print(f"{ {k:v for (k,v) in doc_to_dict(doc).items()} }")

test_dict = doc_to_dict(doc)
for value in test_dict.values():
    print(type(value))
### - Strikethrough on single sentences where keyword(kword) is found is working
### - need to combine runs after this on paragraph so paragraphs can be formatted correctly
''''''
temp = tempfile.NamedTemporaryFile(prefix='test_', 
                                   suffix='.docx', 
                                   dir="./output_files/", 
                                   delete=False)

# output_doc.paragraphs.clear()
# for k, v in test_dict.items():
#     sentence = ". ".join(itertools.chain.from_iterable(v))
#     if "ipsum" in sentence:
#         output_doc.add_paragraph().add_run(f'{sentence}').font.strike = True
#     else:
#         output_doc.add_paragraph().add_run(f"{sentence}")


def style_token(doc: Document, word: str, comment=True) -> Document:
    for p in doc.paragraphs:
        for i, r in enumerate(p.runs):
            if p.runs[i].text.find(word) != -1:
                p.runs[i].font.highlight_color = WD_COLOR_INDEX.RED
                p.runs[i].font.strike = True
    return doc


kword = "ipsum"
for k, list_of_str in test_dict.items():
    print(f"{k}: {list_of_str}")
    print(list_of_str.text.find(kword))
    if list_of_str.text.find(kword) != -1:
        virt_runs = list_of_str.runs
        list_of_str.text = ""
        print(virt_runs)
        for run in virt_runs:
            print(run.text.find(kword))
            if run.text.find(kword) != -1:
                print("Found")
                split_sentence = document_processor.split_sentence_to_list(run.text)
                for word in split_sentence:
                    list_of_str.add_run(word)
                    list_of_str.add_run(" ")


for value in test_dict.values():
    output_doc.add_paragraph(value.text)


style_output_doc = style_token(output_doc, kword, comment=True)

        # if kword in sentence:
        #     print(f"{kword}: FOUND!")
        #     output_doc.add_paragraph().add_run(f'{sentence}').font.strike = True
        # else:
        #     print("Unformatted Sentence")
        #     output_doc.add_paragraph().add_run(f"{sentence}")


# virtual_runs = output_doc.paragraphs
# print(virtual_runs)
style_output_doc.save(temp.name)
''''''





regex = re.compile("ipsum")
def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph


# for k, v in test_dict.items():
#     #print(f'{k}: {v}')
#     for item in v:
#         if "ipsum" in str(item):
#             print(f"{k}: {item}")


# for k, v in test_dict.items():
#     print(f'{k}: {"".join(itertools.chain.from_iterable(v))}')




        # temp_list = extract_words_using_find(v)
        # print(temp_list)
        # output_doc.add_paragraph(v)


# output_doc.save("./output_files/test_doc.docx")





# Can Highlight text
# for par in doc.paragraphs:
#     for run in par.runs:
#         if 'ultricies' in run.text:
#             x = run.text.split('ultricies')
#             run.clear()
#             for i in range(len(x)):
#                 print(x[i])
#                 print("CHANGED")
#                 # run.add_text(x[i])
#                 # run.add_text('ultricies')
#                 # run.font.highlight_color = WD_COLOR_INDEX.YELLOW
#                 # run.font.strike = True

# Remove full paragraphs
# def delete_paragraph(paragraph):
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None

# Create list with single word changed
# for paragraph in doc.paragraphs:
#     if 'lorem' in paragraph.text:
#         for run in paragraph.runs:
#             if 'lorem' in run.text:
#                 x = run.text.split('lorem')
#                 run.clear()
#                 # print(x)
#                 for i in range(len(x)-1):
#                     run.add_text(x[i])
#                     run.add_text('lorem')
#                     run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # new_list.append((par.text).replace("ultricies", "CHANGED"))

# for item in new_list:
#     print(item)

# Clear out working doc
# for par in doc.paragraphs:
#     delete_paragraph(par)

# Fill it with new list of paragraphs with the modifications
# for item in new_list:
#     doc.add_paragraph().add_run(str(item))

# Save it to a new file
# doc.save(output_doc_path)

# @func_log
# def doc_to_list(input_doc: docx.Document):
#     paragraph_lists = [[]]
#     for paragraph in input_doc.paragraphs:
#         paragraph_lists.append(paragraph.text)
#
#     return paragraph_lists


# print(doc_to_list(doc))

# dict_of_para = {}
# for i, paragraph in enumerate(doc.paragraphs):
#     dict_of_para[i] = paragraph.text
#     #print(f"{i}: {paragraph.text}\n")
#
#
# for key, value in dict_of_para.items():
#     print(f"{key}: {value}")


# def doc_to_dict(input_doc: docx.Document) -> dict[int, str]:
#     return {k:v.text for (k,v) in enumerate(input_doc.paragraphs)}
#
#
# test_dict = doc_to_dict(doc)
# for key, value in doc_to_dict(doc).items():
#     print(f"{key}: {value}")
# @func_log
# def test_log_wrapper(x, y):
#     return x+y
#
# result = test_log_wrapper(10, 20)

# NOTES
# import docx

# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)